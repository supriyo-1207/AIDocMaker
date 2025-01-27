from flask import Blueprint, send_file, jsonify
from docx import Document
from bs4 import BeautifulSoup
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from bs4 import NavigableString
import os
import re
from typing import Dict, Tuple, Optional, Any

docs = Blueprint('docs', __name__)

class QuillStyleMapper:
    """Handles Quill-specific style mapping and conversion."""
    
    ALIGNMENT_MAP = {
        'ql-align-center': WD_ALIGN_PARAGRAPH.CENTER,
        'ql-align-right': WD_ALIGN_PARAGRAPH.RIGHT,
        'ql-align-justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        'ql-align-left': WD_ALIGN_PARAGRAPH.LEFT
    }
    
    INDENT_SIZE = 0.5  # inches per indent level
    
    COLOR_PATTERNS = {
        'rgb': re.compile(r'rgb\((\d+),\s*(\d+),\s*(\d+)\)'),
        'hex': re.compile(r'#([0-9a-fA-F]{6}|[0-9a-fA-F]{3})')
    }
    
    @staticmethod
    def parse_color(color_value: str) -> Optional[Tuple[int, int, int]]:
        """Parse color values in RGB or HEX format."""
        if not color_value:
            return None
            
        # Check RGB format
        rgb_match = QuillStyleMapper.COLOR_PATTERNS['rgb'].match(color_value)
        if rgb_match:
            return tuple(map(int, rgb_match.groups()))
        
        # Check HEX format
        hex_match = QuillStyleMapper.COLOR_PATTERNS['hex'].match(color_value)
        if hex_match:
            hex_value = hex_match.group(1)
            if len(hex_value) == 3:
                hex_value = ''.join(c * 2 for c in hex_value)
            return tuple(int(hex_value[i:i+2], 16) for i in (0, 2, 4))
        
        return None

    @staticmethod
    def parse_styles(style_attr: str) -> Dict[str, str]:
        """Parse inline CSS styles into a dictionary."""
        styles = {}
        if not style_attr:
            return styles
            
        for style in style_attr.split(';'):
            if ':' in style:
                key, value = style.split(':', 1)
                styles[key.strip().lower()] = value.strip()
        return styles
    
    @staticmethod
    def get_indent_level(element: Any) -> int:
        """Extract indent level from Quill classes."""
        classes = element.get('class', [])
        if isinstance(classes, str):
            classes = classes.split()
            
        for cls in classes:
            if cls.startswith('ql-indent-'):
                try:
                    return int(cls.replace('ql-indent-', ''))
                except ValueError:
                    pass
        return 0

class DocxFormatter:
    """Handles DOCX document formatting and element processing."""
    
    def __init__(self):
        self.doc = Document()
        self.style_mapper = QuillStyleMapper()
    
    def apply_styles_to_run(self, run: Any, styles: Dict[str, str]) -> None:
        """Apply parsed styles to a DOCX run."""
        if 'color' in styles:
            rgb = self.style_mapper.parse_color(styles['color'])
            if rgb:
                run.font.color.rgb = RGBColor(*rgb)
        
        if 'background-color' in styles:
            rgb = self.style_mapper.parse_color(styles['background-color'])
            if rgb:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        
        if 'font-size' in styles:
            size = styles['font-size']
            if 'px' in size:
                run.font.size = Pt(float(size.replace('px', '')) * 0.75)
        
        if 'font-family' in styles:
            run.font.name = styles['font-family'].strip("'\"")
        
        if 'text-decoration' in styles:
            decoration = styles['text-decoration']
            run.underline = 'underline' in decoration
            run.font.strike = 'line-through' in decoration
    
    def process_element(self, element: Any, paragraph: Any = None) -> None:
        """Process HTML elements and apply formatting."""
        if paragraph is None:
            paragraph = self.doc.add_paragraph()
        
        # Handle indentation
        indent_level = self.style_mapper.get_indent_level(element)
        if indent_level > 0:
            paragraph.paragraph_format.left_indent = Inches(indent_level * self.style_mapper.INDENT_SIZE)
        
        # Handle alignment
        if element.get('class'):
            classes = element['class'] if isinstance(element['class'], list) else element['class'].split()
            for cls, align in self.style_mapper.ALIGNMENT_MAP.items():
                if cls in classes:
                    paragraph.alignment = align
                    break
        
        for child in element.children:
            if isinstance(child, NavigableString):
                if child.strip():
                    run = paragraph.add_run(child.strip())
            else:
                styles = self.style_mapper.parse_styles(child.get('style', ''))
                run = paragraph.add_run(child.get_text(strip=True))
                
                # Apply basic formatting
                if child.name in ['strong', 'b']:
                    run.bold = True
                if child.name in ['em', 'i']:
                    run.italic = True
                if child.name == 'u':
                    run.underline = True
                
                self.apply_styles_to_run(run, styles)
    
    def process_list(self, element: Any, level: int = 0) -> None:
        """Handle ordered and unordered lists with proper formatting."""
        list_style = 'List Bullet' if element.name == 'ul' else 'List Number'
        
        for li in element.find_all('li', recursive=False):
            paragraph = self.doc.add_paragraph(style=list_style)
            paragraph.paragraph_format.left_indent = Inches(level * self.style_mapper.INDENT_SIZE)
            self.process_element(li, paragraph)
            
            # Process nested lists
            nested_lists = li.find_all(['ul', 'ol'], recursive=False)
            for nested_list in nested_lists:
                self.process_list(nested_list, level + 1)
    
    def convert_html_to_docx(self, html_content: str) -> str:
        """Convert HTML content to DOCX format."""
        soup = BeautifulSoup(html_content, 'html.parser')
        
        for element in soup.children:
            if isinstance(element, NavigableString):
                if element.strip():
                    self.process_element(element)
            else:
                if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    paragraph = self.doc.add_paragraph(style=f'Heading {element.name[1]}')
                    self.process_element(element, paragraph)
                elif element.name in ['ul', 'ol']:
                    self.process_list(element)
                else:  # p, div, etc.
                    self.process_element(element)
        
        os.makedirs('outputs', exist_ok=True)
        file_path = os.path.join('outputs', 'generated_doc.docx')
        self.doc.save(file_path)
        return file_path

@docs.route('/generate-docx', methods=['GET'])
def generate_docx():
    from controllers.page_controller import result_manager
    
    content_to_use = result_manager.get_result()
    if not content_to_use:
        return jsonify({"error": "No content available"}), 400
    
    try:
        formatter = DocxFormatter()
        file_path = formatter.convert_html_to_docx(content_to_use)
        return send_file(file_path, as_attachment=True, download_name='generated_report.docx')
    except Exception as e:
        return jsonify({"error": str(e)}), 500