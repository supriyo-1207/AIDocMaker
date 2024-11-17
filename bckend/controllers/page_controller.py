from flask import Blueprint, jsonify, request, send_file
from langchain.prompts import PromptTemplate
from langchain_google_genai import ChatGoogleGenerativeAI
import os
from dotenv import load_dotenv
import markdown2
from docx import Document
from bs4 import BeautifulSoup

# Load environment variables
load_dotenv()

# Create a Blueprint instance
page = Blueprint('page', __name__)

# Initialize the LLM model
os.environ["GOOGLE_API_KEY"] = os.getenv("GOOGLE_API_KEY")
llm = ChatGoogleGenerativeAI(model="gemini-pro")

# Define the report template
report_template = """Generate a professional report on the following topic:
Topic: {topic}

Requirements:
1. Provide an introduction.
2. Add key details with bullet points.
3. Conclude with actionable insights.
"""

# Create a Prompt Template
prompt = PromptTemplate(
    input_variables=["topic"],
    template=report_template,
)

# Use RunnableSequence syntax by combining prompt and llm
analysis_chain = prompt | llm
result = {}

# Define the route
@page.route('/getdata', methods=['POST'])
def gemini():
    data = request.get_json()
    print(data)
    
    # Get the topic and background from the request
    topic = data.get('instructions')
    background = data.get('background')

    # Check if the topic is provided
    if not topic:
        return jsonify({"error": "Topic is required"}), 400
    
    # Append background information to the topic if provided
    if background:
        topic += f"\n\nBackground: {background}"
    
    # Generate the output by invoking the chain
    output = analysis_chain.invoke({"topic": topic})
    global result

    # Extract the generated text from the AIMessage object
    generated_text = output.content  # Accessing the content attribute

    # Convert the generated text to markdown
    result = markdown2.markdown(generated_text)
    print(result)
    
    # Return the generated report
    return jsonify({"report": result}), 200

@page.route('/getdata', methods=['GET'])
def getdata():
    print(result)
    return jsonify({"report": result}), 200


# Your route for generating docx
@page.route('/generate-docx', methods=['GET'])
def generate_docx():
    # Function to generate the docx from HTML content
    def generate_docx_from_html(html_content):
        # Create a new Document
        doc = Document()

        # Parse the HTML content using BeautifulSoup
        soup = BeautifulSoup(html_content, 'html.parser')

        # Iterate through the parsed HTML and handle different tags
        for element in soup.find_all(True):  # Find all tags
            if element.name == 'h1':
                # For <h1>, make it a Heading 1
                doc.add_paragraph(element.get_text(), style='Heading 1')
            elif element.name == 'h2':
                # For <h2>, make it a Heading 2
                doc.add_paragraph(element.get_text(), style='Heading 2')
            elif element.name == 'h3':
                # For <h3>, make it a Heading 3
                doc.add_paragraph(element.get_text(), style='Heading 3')
            elif element.name == 'ul':
                # For <ul>, create bullet points
                for li in element.find_all('li'):
                    doc.add_paragraph(li.get_text(), style='List Bullet')
            elif element.name == 'ol':
                # For <ol>, create numbered list (ordered list)
                for li in element.find_all('li'):
                    doc.add_paragraph(li.get_text(), style='List Number')
            elif element.name == 'strong':
                # For <strong>, make the text bold
                para = doc.add_paragraph()
                run = para.add_run(element.get_text())
                run.bold = True
            elif element.name == 'em':
                # For <em>, make the text italic
                para = doc.add_paragraph()
                run = para.add_run(element.get_text())
                run.italic = True
            elif element.name == 'u':
                # For <u>, underline the text
                para = doc.add_paragraph()
                run = para.add_run(element.get_text())
                run.underline = True
            else:
                # For other tags, just add text as a regular paragraph
                doc.add_paragraph(element.get_text())

        # Save the document to a temporary file
        docx_file = 'generated_doc.docx'
        doc.save(docx_file)

        # Return the document file name
        return docx_file

    # Generate the docx file from the result (HTML content)
    docx_file = generate_docx_from_html(result)

    # Send the file as a response
    return send_file(docx_file, as_attachment=True)
